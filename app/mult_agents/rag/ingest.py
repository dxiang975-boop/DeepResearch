import json
import logging
import os
import sys
from pathlib import Path
from typing import Callable


# 将项目根目录添加到 PYTHONPATH，解决直接运行脚本时的模块导入问题。
project_root = Path(__file__).resolve().parents[3]
if str(project_root) not in sys.path:
    sys.path.insert(0, str(project_root))

# 先加载 .env，再导入配置模块，确保 Milvus、DashScope 等配置可用。
from dotenv import load_dotenv

env_path = project_root / ".env"
if env_path.exists():
    load_dotenv(dotenv_path=env_path)

from mult_agents.config import AppConfig
from mult_agents.rag.core import RAGConfig, RAGSystem


INPUT_PATH = Path(os.getenv("RAG_INPUT_PATH", "knowledge_base"))
COLLECTION_NAME = ""
MILVUS_HOST = ""
MILVUS_PORT = 0
EMBEDDING_MODEL = "text-embedding-v1"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

SUPPORTED_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".log",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
    ".tsv",
    ".xlsx",
    ".xls",
    ".pdf",
    ".docx",
    ".py",
}


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_json(path: Path) -> str:
    data = json.loads(path.read_text(encoding="utf-8"))
    return json.dumps(data, ensure_ascii=False, indent=2)


def _read_table(path: Path) -> str:
    import pandas as pd

    if path.suffix.lower() == ".tsv":
        df = pd.read_csv(path, sep="\t")
    else:
        df = pd.read_csv(path)
    return df.to_csv(index=False)


def _read_excel(path: Path) -> str:
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None)
    parts: list[str] = []
    for name, df in sheets.items():
        parts.append(f"# Sheet: {name}")
        parts.append(df.to_csv(index=False))
    return "\n\n".join(parts)


def _resolve_input_path(input_path: Path) -> Path:
    """将相对路径固定解析到项目根目录，避免从不同工作目录运行时读错文件夹。"""
    expanded = input_path.expanduser()
    if expanded.is_absolute():
        return expanded.resolve()
    return (project_root / expanded).resolve()


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("读取 PDF 需要安装可选依赖：pip install pypdf") from exc

    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"# Page {index}\n{text}")
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("读取 DOCX 需要安装可选依赖：pip install python-docx") from exc

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


READERS: dict[str, Callable[[Path], str]] = {
    ".txt": _read_text_file,
    ".md": _read_text_file,
    ".markdown": _read_text_file,
    ".log": _read_text_file,
    ".yaml": _read_text_file,
    ".yml": _read_text_file,
    ".json": _read_json,
    ".csv": _read_table,
    ".tsv": _read_table,
    ".xlsx": _read_excel,
    ".xls": _read_excel,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".py": _read_text_file,
}


def _collect_paths(input_path: Path) -> list[Path]:
    if input_path.is_file():
        return [input_path] if input_path.suffix.lower() in SUPPORTED_SUFFIXES else []

    paths: list[Path] = []
    for suffix in SUPPORTED_SUFFIXES:
        paths.extend(sorted(input_path.rglob(f"*{suffix}")))
    return sorted(set(paths))


def _doc_type(path: Path) -> str:
    suffix = path.suffix.lower().lstrip(".")
    if suffix in {"pdf", "docx"}:
        return "paper_or_document"
    if suffix in {"csv", "tsv", "xlsx", "xls"}:
        return "experiment_table"
    if suffix in {"log"}:
        return "experiment_log"
    if suffix in {"py"}:
        return "python_code"
    if suffix in {"md", "markdown", "txt"}:
        return "text_note"
    return suffix or "unknown"


def _metadata_for_path(path: Path) -> dict:
    return {
        "source": str(path),
        "title": path.stem,
        "file_name": path.name,
        "file_suffix": path.suffix.lower(),
        "doc_type": _doc_type(path),
    }


def _read_document(path: Path) -> str:
    reader = READERS.get(path.suffix.lower())
    if reader is None:
        raise ValueError(f"不支持的文件类型：{path.suffix}")
    text = reader(path)
    if not text.strip():
        raise ValueError(f"文件没有提取到可入库文本：{path}")
    return text


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s | %(levelname)s | %(message)s")

    config = AppConfig.from_file()
    collection_name = COLLECTION_NAME or config.milvus_collection
    milvus_host = MILVUS_HOST or config.milvus_host
    milvus_port = MILVUS_PORT or config.milvus_port
    rag_cfg = RAGConfig(
        milvus_host=milvus_host,
        milvus_port=milvus_port,
        collection_name=collection_name,
        embedding_model=EMBEDDING_MODEL,
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    rag = RAGSystem(api_key=config.api_key, config=rag_cfg)

    input_path = _resolve_input_path(INPUT_PATH)
    if not input_path.exists():
        raise FileNotFoundError(f"未找到入库路径：{input_path}")

    paths = _collect_paths(input_path)
    if not paths:
        raise ValueError(f"未找到可入库文件：{input_path}")

    total_chunks = 0
    skipped: list[str] = []
    for path in paths:
        try:
            text = _read_document(path)
            total_chunks += rag.ingest_text(text, source=str(path), metadata=_metadata_for_path(path))
            logging.info("入库完成 | file=%s", path)
        except Exception as exc:
            skipped.append(f"{path}: {exc}")
            logging.warning("跳过文件 | file=%s | reason=%s", path, exc)

    print(f"入库完成 | 文件数={len(paths)} | chunk数={total_chunks} | collection={collection_name}")
    if skipped:
        print("以下文件未入库：")
        for item in skipped:
            print(f"- {item}")


if __name__ == "__main__":
    main()
